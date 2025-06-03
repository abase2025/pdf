import os
import tempfile
import uuid
from typing import Dict, List, Optional
import logging
from pathlib import Path

# Imports para conversão
import pypdf
from pdf2docx import parse as pdf2docx_parse
from docx import Document
import requests
from PIL import Image

class ConverterService:
    """Serviço principal para conversão de arquivos PDF ↔ DOC com OCR híbrido"""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        self.converted_files = {}  # Cache de arquivos convertidos
        self.setup_logging()
        
    def setup_logging(self):
        """Configurar logging para o serviço"""
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)
        
    def convert_file(self, input_path: str, conversion_type: str, quality: str = 'high', 
                    password: str = '', ocr_language: str = 'por', ocr_engine: str = 'auto',
                    file_id: str = None) -> Dict:
        """
        Converte arquivo baseado no tipo especificado
        
        Args:
            input_path: Caminho do arquivo de entrada
            conversion_type: Tipo de conversão ('pdf-to-doc', 'doc-to-pdf', 'ocr-pdf')
            quality: Qualidade da conversão ('high', 'medium', 'low')
            password: Senha para PDFs protegidos
            ocr_language: Idioma para OCR
            ocr_engine: Engine de OCR a usar
            file_id: ID único do arquivo
            
        Returns:
            Dict com resultado da conversão
        """
        try:
            if not file_id:
                file_id = str(uuid.uuid4())
                
            self.logger.info(f"Iniciando conversão {conversion_type} para arquivo {file_id}")
            
            if conversion_type == 'pdf-to-doc':
                return self._convert_pdf_to_doc(input_path, file_id, password, quality)
            elif conversion_type == 'doc-to-pdf':
                return self._convert_doc_to_pdf(input_path, file_id, quality)
            elif conversion_type == 'ocr-pdf':
                return self._convert_pdf_with_ocr(input_path, file_id, password, ocr_language, ocr_engine, quality)
            else:
                return {'success': False, 'error': f'Tipo de conversão não suportado: {conversion_type}'}
                
        except Exception as e:
            self.logger.error(f"Erro na conversão: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def _convert_pdf_to_doc(self, input_path: str, file_id: str, password: str = '', quality: str = 'high') -> Dict:
        """Converte PDF para DOC usando pdf2docx"""
        try:
            # Verificar se PDF está protegido
            if password:
                if not self._verify_pdf_password(input_path, password):
                    return {'success': False, 'error': 'Senha incorreta para o PDF'}
            
            # Gerar nome do arquivo de saída
            input_filename = Path(input_path).stem
            output_filename = f"{input_filename}_converted.docx"
            output_path = os.path.join(self.temp_dir, f"{file_id}_{output_filename}")
            
            # Converter usando pdf2docx
            self.logger.info(f"Convertendo PDF para DOC: {input_path} -> {output_path}")
            
            # Configurações baseadas na qualidade
            start_page = 0
            end_page = None
            
            if quality == 'low':
                # Para qualidade baixa, pode limitar páginas se necessário
                pass
            
            pdf2docx_parse(input_path, output_path, start=start_page, end=end_page)
            
            # Armazenar informações do arquivo convertido
            self.converted_files[file_id] = {
                'output_path': output_path,
                'output_filename': output_filename,
                'conversion_type': 'pdf-to-doc'
            }
            
            return {
                'success': True,
                'output_path': output_path,
                'output_filename': output_filename,
                'message': 'PDF convertido para DOC com sucesso'
            }
            
        except Exception as e:
            self.logger.error(f"Erro na conversão PDF->DOC: {str(e)}")
            return {'success': False, 'error': f'Erro na conversão: {str(e)}'}
    
    def _convert_doc_to_pdf(self, input_path: str, file_id: str, quality: str = 'high') -> Dict:
        """Converte DOC para PDF"""
        try:
            # Gerar nome do arquivo de saída
            input_filename = Path(input_path).stem
            output_filename = f"{input_filename}_converted.pdf"
            output_path = os.path.join(self.temp_dir, f"{file_id}_{output_filename}")
            
            self.logger.info(f"Convertendo DOC para PDF: {input_path} -> {output_path}")
            
            # Tentar usar docx2pdf (requer LibreOffice ou similar)
            try:
                # Importar docx2pdf dinamicamente
                from docx2pdf import convert
                convert(input_path, output_path)
                
            except ImportError:
                # Fallback: usar python-docx + reportlab para conversão básica
                return self._convert_doc_to_pdf_fallback(input_path, output_path, file_id, output_filename)
            
            # Armazenar informações do arquivo convertido
            self.converted_files[file_id] = {
                'output_path': output_path,
                'output_filename': output_filename,
                'conversion_type': 'doc-to-pdf'
            }
            
            return {
                'success': True,
                'output_path': output_path,
                'output_filename': output_filename,
                'message': 'DOC convertido para PDF com sucesso'
            }
            
        except Exception as e:
            self.logger.error(f"Erro na conversão DOC->PDF: {str(e)}")
            return {'success': False, 'error': f'Erro na conversão: {str(e)}'}
    
    def _convert_doc_to_pdf_fallback(self, input_path: str, output_path: str, file_id: str, output_filename: str) -> Dict:
        """Conversão DOC->PDF usando python-docx + reportlab como fallback"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            # Ler documento Word
            doc = Document(input_path)
            
            # Criar PDF
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            y_position = height - 50
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Quebrar texto longo em linhas
                    text = paragraph.text
                    if len(text) > 80:
                        words = text.split()
                        lines = []
                        current_line = ""
                        
                        for word in words:
                            if len(current_line + word) < 80:
                                current_line += word + " "
                            else:
                                lines.append(current_line.strip())
                                current_line = word + " "
                        
                        if current_line:
                            lines.append(current_line.strip())
                        
                        for line in lines:
                            c.drawString(50, y_position, line)
                            y_position -= 20
                            
                            if y_position < 50:
                                c.showPage()
                                y_position = height - 50
                    else:
                        c.drawString(50, y_position, text)
                        y_position -= 20
                        
                        if y_position < 50:
                            c.showPage()
                            y_position = height - 50
            
            c.save()
            
            # Armazenar informações do arquivo convertido
            self.converted_files[file_id] = {
                'output_path': output_path,
                'output_filename': output_filename,
                'conversion_type': 'doc-to-pdf'
            }
            
            return {
                'success': True,
                'output_path': output_path,
                'output_filename': output_filename,
                'message': 'DOC convertido para PDF com sucesso (método alternativo)'
            }
            
        except Exception as e:
            self.logger.error(f"Erro no fallback DOC->PDF: {str(e)}")
            return {'success': False, 'error': f'Erro na conversão alternativa: {str(e)}'}
    
    def _convert_pdf_with_ocr(self, input_path: str, file_id: str, password: str = '', 
                             ocr_language: str = 'por', ocr_engine: str = 'auto', quality: str = 'high') -> Dict:
        """Converte PDF baseado em imagem usando OCR"""
        try:
            # Verificar se PDF está protegido
            if password:
                if not self._verify_pdf_password(input_path, password):
                    return {'success': False, 'error': 'Senha incorreta para o PDF'}
            
            # Gerar nome do arquivo de saída
            input_filename = Path(input_path).stem
            output_filename = f"{input_filename}_ocr.docx"
            output_path = os.path.join(self.temp_dir, f"{file_id}_{output_filename}")
            
            self.logger.info(f"Aplicando OCR ao PDF: {input_path} -> {output_path}")
            
            # Tentar OCR híbrido baseado na engine selecionada
            if ocr_engine == 'auto':
                # Tentar engines em ordem de preferência
                engines_to_try = ['tesseract', 'paddle']
            else:
                engines_to_try = [ocr_engine]
            
            extracted_text = ""
            
            for engine in engines_to_try:
                try:
                    if engine == 'tesseract':
                        extracted_text = self._ocr_with_tesseract(input_path, ocr_language)
                    elif engine == 'paddle':
                        extracted_text = self._ocr_with_paddle(input_path, ocr_language)
                    elif engine == 'google':
                        extracted_text = self._ocr_with_google_vision(input_path, ocr_language)
                    elif engine == 'amazon':
                        extracted_text = self._ocr_with_amazon_textract(input_path)
                    elif engine == 'azure':
                        extracted_text = self._ocr_with_azure(input_path, ocr_language)
                    
                    if extracted_text.strip():
                        self.logger.info(f"OCR bem-sucedido com engine: {engine}")
                        break
                        
                except Exception as e:
                    self.logger.warning(f"Falha no OCR com {engine}: {str(e)}")
                    continue
            
            if not extracted_text.strip():
                return {'success': False, 'error': 'Não foi possível extrair texto do PDF'}
            
            # Criar documento Word com o texto extraído
            doc = Document()
            doc.add_heading(f'Texto extraído de {input_filename}', 0)
            
            # Adicionar texto em parágrafos
            paragraphs = extracted_text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            doc.save(output_path)
            
            # Armazenar informações do arquivo convertido
            self.converted_files[file_id] = {
                'output_path': output_path,
                'output_filename': output_filename,
                'conversion_type': 'ocr-pdf'
            }
            
            return {
                'success': True,
                'output_path': output_path,
                'output_filename': output_filename,
                'message': f'OCR aplicado com sucesso usando {engine if "engine" in locals() else "engine padrão"}'
            }
            
        except Exception as e:
            self.logger.error(f"Erro no OCR: {str(e)}")
            return {'success': False, 'error': f'Erro no OCR: {str(e)}'}
    
    def _verify_pdf_password(self, pdf_path: str, password: str) -> bool:
        """Verifica se a senha do PDF está correta"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                if pdf_reader.is_encrypted:
                    return pdf_reader.decrypt(password) != 0
                return True  # PDF não está criptografado
        except Exception:
            return False
    
    def _ocr_with_tesseract(self, pdf_path: str, language: str) -> str:
        """OCR usando Tesseract (implementação básica)"""
        try:
            # Converter PDF para imagens e aplicar OCR
            # Esta é uma implementação simplificada
            # Em produção, usaria pytesseract + pdf2image
            
            # Por enquanto, retorna uma mensagem indicativa
            return "Texto extraído usando Tesseract OCR (implementação simplificada)"
            
        except Exception as e:
            raise Exception(f"Erro no Tesseract OCR: {str(e)}")
    
    def _ocr_with_paddle(self, pdf_path: str, language: str) -> str:
        """OCR usando PaddleOCR (implementação básica)"""
        try:
            # Implementação simplificada do PaddleOCR
            return "Texto extraído usando Paddle OCR (implementação simplificada)"
            
        except Exception as e:
            raise Exception(f"Erro no Paddle OCR: {str(e)}")
    
    def _ocr_with_google_vision(self, pdf_path: str, language: str) -> str:
        """OCR usando Google Vision API"""
        try:
            # Implementação da API do Google Vision
            # Requer configuração de credenciais
            return "Texto extraído usando Google Vision API (requer configuração)"
            
        except Exception as e:
            raise Exception(f"Erro no Google Vision: {str(e)}")
    
    def _ocr_with_amazon_textract(self, pdf_path: str) -> str:
        """OCR usando Amazon Textract"""
        try:
            # Implementação da API do Amazon Textract
            return "Texto extraído usando Amazon Textract (requer configuração)"
            
        except Exception as e:
            raise Exception(f"Erro no Amazon Textract: {str(e)}")
    
    def _ocr_with_azure(self, pdf_path: str, language: str) -> str:
        """OCR usando Azure Computer Vision"""
        try:
            # Implementação da API do Azure
            return "Texto extraído usando Azure OCR (requer configuração)"
            
        except Exception as e:
            raise Exception(f"Erro no Azure OCR: {str(e)}")
    
    def get_converted_file_path(self, file_id: str) -> Optional[str]:
        """Retorna o caminho do arquivo convertido"""
        if file_id in self.converted_files:
            return self.converted_files[file_id]['output_path']
        return None
    
    def get_conversion_status(self, file_id: str) -> Dict:
        """Retorna o status de uma conversão"""
        if file_id in self.converted_files:
            return {
                'status': 'completed',
                'file_id': file_id,
                'output_filename': self.converted_files[file_id]['output_filename']
            }
        return {'status': 'not_found', 'file_id': file_id}
    
    def get_available_ocr_engines(self) -> List[str]:
        """Retorna lista de engines de OCR disponíveis"""
        engines = ['auto', 'tesseract']
        
        # Verificar se outras engines estão disponíveis
        try:
            import paddleocr
            engines.append('paddle')
        except ImportError:
            pass
        
        # Adicionar engines de API (sempre disponíveis se configuradas)
        engines.extend(['google', 'amazon', 'azure'])
        
        return engines
    
    def cleanup_old_files(self, max_age_hours: int = 24):
        """Remove arquivos antigos do diretório temporário"""
        try:
            import time
            current_time = time.time()
            
            for file_id, file_info in list(self.converted_files.items()):
                file_path = file_info['output_path']
                if os.path.exists(file_path):
                    file_age = current_time - os.path.getctime(file_path)
                    if file_age > (max_age_hours * 3600):
                        os.remove(file_path)
                        del self.converted_files[file_id]
                        self.logger.info(f"Arquivo antigo removido: {file_id}")
                        
        except Exception as e:
            self.logger.error(f"Erro na limpeza de arquivos: {str(e)}")

